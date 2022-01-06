from django.http.response import Http404
from django.shortcuts import render,redirect
from django.contrib.auth import login,authenticate
from .models import tbl_Authentication
from django.http import HttpResponse, JsonResponse
from wsgiref.util import FileWrapper
from .models import upload_documents

# documentation: https://pdfminersix.readthedocs.io/en/latest/index.html
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser

from django.conf import settings


import os
import nltk
import copy

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from textblob import TextBlob
from textblob import classifiers

# Create your views here.
parsed_text = {}
parsed_text['parsed_text'] = ''
parsed_text['doc_name'] = ''

copy_file = ''

#4
training_set = []
train_directory="../productsite/documents/trainingsets/"

def login(request):
    return render(request, 'login.html')

def user_login(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        try:
            user = True
            #user = tbl_Authentication.empAuth_objects.get(username=username,password=password)
            if user is not None:
                return render(request, 'dashboard.html', parsed_text)
            else:
                print("Someone tried to login and failed.")
                print("They used username: {} and password: {}".format(username,password))
                return redirect('/')
        except Exception as identifier:
            return redirect('/')
    else:
        return render(request,'login.html')

def file_upload_view(request):
    # print(request.FILES)
    if request.method == 'POST':
        my_file = request.FILES.get('file')
        upload_documents.objects.create(upload=my_file)
        file_name_despaced = my_file.name.replace(" ", "_")
       
        #Might be removed.
        if 'pdf' in file_name_despaced: 
            parsed_text['parsed_text'] = parse_uploaded_pdf(file_name_despaced)
        else:
            parsed_text['parsed_text'] = parse_uploaded_docx(file_name_despaced)
                
        parsed_text['doc_name'] = ' - ' + file_name_despaced

        #Creating dictionary of dictionaries, and dictionary of {paragraph:results}
        document_content = directory_docx_content(train_directory)
        textblob_dictionary = doc_polarity(document_content)
        textblob_dictionary = polarity_label(textblob_dictionary)

        #5
        #Classification, and training_set 
        training_set = list(textblob_dictionary.items())
        classifier = classifiers.DecisionTreeClassifier(training_set)

        #Display document with highlighted text.
        #parsed_text['parsed_text'] = parse_uploaded_docx(highlight_document(file_name_despaced, classifier, textblob_dictionary))
        # parsed_text['parsed_text'] = highlight_document(file_name_despaced, classifier, textblob_dictionary)
        test_dictionary = docx_to_dict(f"../productsite/documents/contracts/" + file_name_despaced)
        filename = file_name_despaced.split('.')[0]
        pos_score, neg_score, pos_paragraph, neg_paragraph = highlight_document(filename, classifier, test_dictionary)
        copy_file = file_name_despaced.split('.')[0] + "_copy.docx"

        parsed_text['pos_paragraph'] = pos_paragraph
        parsed_text['neg_paragraph'] = neg_paragraph

        # wont retain colors since it just extracts text into a string. probably need to loop a dictionary with neutral, pos, neg and change the font color with html for all pos, neg
        # parsed_text['parsed_text'] = parse_uploaded_docx(copy_file)
        #Display accuracy
        parsed_text['accuracy'] = classifier_accuracy(file_name_despaced, classifier)
        
        total = pos_score + neg_score
        pos_percent = pos_score / total
        neg_percent = neg_score / total
        
        #Display the positive and negative accuracy.
        parsed_text['favourable'] = pos_percent
        parsed_text['unfavourable'] = neg_percent
                
        return JsonResponse({'parsed_text': parsed_text['parsed_text']})
    return JsonResponse({'post':'false'})

def download(request, copy_file, path='contract/'):
    file_path=os.path.join(settings.MEDIA_ROOT,path, copy_file)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response=HttpResponse(fh.read(), content_type="application/docx")
            response['Content-Disposition']='inline;filename='+os.path.basename(file_path)
            return response
    raise Http404

# method for pdf text extraction
def parse_uploaded_pdf(served):
    output_string = StringIO()

    # should be path and name of uploaded doc
    filename = f"../productsite/documents/contracts/{served}"
    with open(filename, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
        output_string = output_string.getvalue().replace('\n', '')
        return output_string

# method for word doc text extraction
def parse_uploaded_docx(served):
    docx_dict = {}
    document = Document(f"../productsite/documents/contracts/{served}")
    i = 0
    for p in document.paragraphs:
        i+=1
        if (len(p.text)>0):
            docx_dict[i] = p.text
    document_text = ','.join(docx_dict.values())
    return document_text

# method for extraction of training set
def docx_to_dict(file):
    docx_dict = {}
    document = Document(file)
    i = 0
    for p in document.paragraphs:
        i+=1
        if (len(p.text)>0):
            docx_dict[i] = p.text
    return docx_dict

#1
# method for creating a dictionary of dictionaries containing 
# content of all files in directory
def directory_docx_content(directory):
    docx_content = {}
    for file in os.listdir(directory):
        if file.endswith(".docx") or file.endswith(".doc"):
            path_to_docx = os.path.join(directory, file)
            docx_content[path_to_docx] = docx_to_dict(path_to_docx)
    return docx_content

#2
# method for joining the content of all the training folder 
# documents to create a dictionary of {paragraphs:results}
def doc_polarity(docx_content):
    textblob_dict = {}
    for file,content in docx_content.items():
        for number,text in content.items():
            polarity_result = TextBlob(text).sentiment.polarity
            if polarity_result != 0:
                textblob_dict.update({text:polarity_result})
    return textblob_dict

# method for accuracy metric of single test document
def test_polarity(docx_content):
    textblob_dict = {}
    for number,text in docx_content.items():
        polarity_result = TextBlob(text).sentiment.polarity
        if polarity_result != 0:
                textblob_dict.update({text:polarity_result})
    return textblob_dict

#3
# converting sentiment value into a string
# NOTE: 0.3 and -0.1 are arbitrarily chosen and can be modified
def polarity_label(textblob_dict):
    for k,v in textblob_dict.items():
        if v > 0.3:
            textblob_dict[k] = 'positive'
        elif v < -0.1:
            textblob_dict[k] = 'negative'
        else:
            textblob_dict[k] = 'neutral'
    return textblob_dict

#note 
def dict_to_list(textblob_dict):
    return list(textblob_dict.items())
    
#6
def highlight_document(file, classifier, dictionary):
    pos_counter = 0
    neg_counter = 0
    pos_paragraph = {}
    neg_paragraph = {}
    doc = Document()
    filename = f"{file}_copy.docx"
    filepath = f"../productsite/documents/contracts/" + filename
    for number, text in dictionary.items():
        blob = TextBlob(text, classifier=classifier)
        if(blob.classify() == 'positive'): 
            pos_paragraph.update({number:text})
            pos_counter += 1
            run = doc.add_paragraph().add_run(text)
            font = run.font
            font.color.rgb = RGBColor(34, 204, 0)
            doc.save(filepath)
        elif(blob.classify() == 'negative'):
            neg_paragraph.update({number:text})
            neg_counter += 1
            run = doc.add_paragraph().add_run(text)
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)
            doc.save(filepath)
        else:
            doc.add_paragraph(text)
            doc.save(filepath)
    return pos_counter, neg_counter, pos_paragraph, neg_paragraph
#7
# method for getting accuracy of classifier
def classifier_accuracy(file, classifier):
    test_content = docx_to_dict(f"../productsite/documents/contracts/" + file)
    testing_label = test_polarity(test_content)
    testing_label = polarity_label(testing_label)
    test_set = []
    test_set = list(testing_label.items())
    return (classifier.accuracy(test_set))
    

